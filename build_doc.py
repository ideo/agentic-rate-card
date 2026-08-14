from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("AI_Project_Token_Cost_Guide.docx")

INK = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MID_BLUE = "557A95"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "202B33"
WHITE = "FFFFFF"
BORDER = "C9D3DD"
GOLD = "A26B16"

MODEL_PRICES = {
    "Luna": (0.20, 1.20),
    "Terra": (2.00, 12.00),
    "Sol": (5.00, 30.00),
    "Haiku 4.5": (1.00, 5.00),
    "Sonnet 5": (2.00, 10.00),
    "Opus 5": (5.00, 25.00),
    "Fable 5": (10.00, 50.00),
}


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size="6"):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def keep_table_row_together(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = trPr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")

    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_font(run, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8, color=MID_GRAY)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    keepNext = pPr.find(qn("w:keepNext"))
    if keepNext is None:
        keepNext = OxmlElement("w:keepNext")
        pPr.append(keepNext)


def style_paragraph(paragraph, before=0, after=6, line=1.25, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_para(doc, text="", size=11, color=DARK, bold=False, italic=False,
             before=0, after=6, line=1.25, align=None, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, before, after, line, align)
    if keep:
        keep_with_next(p)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_label_para(doc, label, text, after=6):
    p = doc.add_paragraph()
    style_paragraph(p, 0, after, 1.25)
    r = p.add_run(label + " ")
    set_font(r, size=11, bold=True, color=INK)
    r = p.add_run(text)
    set_font(r, size=11, color=DARK)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    style_paragraph(p, 4, 10, 1.2)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.08)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_BLUE)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(label + "  ")
    set_font(r, size=10.5, bold=True, color=INK)
    r = p.add_run(text)
    set_font(r, size=10.5, color=DARK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    keep_with_next(p)
    return p


def add_section_kicker(doc, text):
    p = add_para(doc, text.upper(), size=9.5, color=GOLD, bold=True,
                 before=0, after=4, line=1.0, keep=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_caption(doc, text):
    return add_para(doc, text, size=8.5, color=MID_GRAY, italic=True,
                    before=4, after=5, line=1.1)


def money(value):
    if value < 0.001:
        return "<$0.001"
    if value < 0.01:
        return f"${value:.3f}".rstrip("0").rstrip(".")
    if value < 1:
        return f"${value:.2f}"
    if value < 100:
        return f"${value:.1f}" if abs(value - round(value)) > 0.001 else f"${value:.0f}"
    return f"${value:,.0f}"


def task_cost(model, input_tokens, output_tokens):
    i, o = MODEL_PRICES[model]
    return input_tokens / 1_000_000 * i + output_tokens / 1_000_000 * o


def cost_range(model, lo_i, hi_i, lo_o, hi_o):
    return f"{money(task_cost(model, lo_i, lo_o))}–{money(task_cost(model, hi_i, hi_o))}"


def fmt_tokens(lo_i, hi_i, lo_o, hi_o):
    def val(n):
        if n >= 1_000_000:
            x = n / 1_000_000
            return f"{x:g}M"
        return f"{n / 1000:g}K"
    return f"{val(lo_i)}–{val(hi_i)} in\n{val(lo_o)}–{val(hi_o)} out"


def write_cell(cell, lines, header=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=8.4,
               bold_first=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    values = lines if isinstance(lines, list) else [lines]
    for idx, line in enumerate(values):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1 if idx == 0 else 0)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(line)
        set_font(r, size=size, bold=header or (bold_first and idx == 0),
                 color=WHITE if header else DARK)


def add_matrix(doc, rows):
    headers = ["Use case", "What the work includes", "Cumulative tokens", "OpenAI", "Anthropic"]
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [1700, 3300, 1350, 1505, 1505]
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        set_cell_shading(hdr.cells[idx], INK)
        write_cell(hdr.cells[idx], text, header=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER if idx >= 2 else WD_ALIGN_PARAGRAPH.LEFT,
                   size=8.3)
    for ridx, row in enumerate(rows):
        tr = table.add_row()
        keep_table_row_together(tr)
        if ridx % 2 == 1:
            for c in tr.cells:
                set_cell_shading(c, PALE_BLUE)
        usecase, desc, li, hi, lo, ho, om, am = row
        write_cell(tr.cells[0], usecase, size=8.5, bold_first=True)
        write_cell(tr.cells[1], desc, size=8.3)
        write_cell(tr.cells[2], fmt_tokens(li, hi, lo, ho).split("\n"),
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=8.2, bold_first=True)
        write_cell(tr.cells[3], [om, cost_range(om, li, hi, lo, ho)],
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=8.2, bold_first=True)
        write_cell(tr.cells[4], [am, cost_range(am, li, hi, lo, ho)],
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=8.2, bold_first=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_simple_table(doc, headers, rows, widths, aligns=None, font_size=9.0, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        set_cell_shading(hdr.cells[idx], INK)
        write_cell(hdr.cells[idx], text, header=True,
                   align=(aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT), size=font_size)
    for ridx, row in enumerate(rows):
        tr = table.add_row()
        keep_table_row_together(tr)
        if zebra and ridx % 2 == 1:
            for c in tr.cells:
                set_cell_shading(c, PALE_BLUE)
        for idx, value in enumerate(row):
            lines = value if isinstance(value, list) else [value]
            write_cell(tr.cells[idx], lines,
                       align=(aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT),
                       size=font_size,
                       bold_first=(idx == 0 or len(lines) > 1))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True


def setup_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AI PROJECT EFFORT · TOKEN & COST GUIDE")
    set_font(r, size=8, bold=True, color=MID_GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def add_cover(doc):
    add_para(doc, "PLANNING REFERENCE", size=10, color=GOLD, bold=True,
             before=76, after=16, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "AI Project Effort,\nToken & Cost Guide", size=28, color=INK, bold=True,
             before=0, after=12, line=1.02, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "A practical matrix for everyday tasks, connected tools,\nsoftware builds, parallel agents, and deep knowledge work",
             size=14, color=MID_BLUE, before=0, after=26, line=1.2,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "API-equivalent estimates · USD · pricing checked August 12, 2026",
             size=10, color=MID_GRAY, italic=True, before=0, after=38, line=1.0,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "THE SHORT VERSION",
        "A message or single tool action is usually fractions of a cent to a few cents. "
        "A strong synthesis, document, or code change is often cents to low dollars. "
        "A full day of frontier-model coding or deep knowledge work is commonly tens of dollars in raw model tokens; "
        "a multi-agent build can reach the low hundreds. Human time, tool fees, hosting, and iteration usually matter more than tokens.",
    )
    add_heading(doc, "What this guide is—and is not", 2)
    add_label_para(doc, "It is", "a budgeting framework for cumulative API token use across all model calls in a completed task.")
    add_label_para(doc, "It is not", "a quote, a subscription-plan comparison, or a promise that every project will fit the band shown.")
    add_label_para(doc, "Best use", "pick the nearest use case, choose a model tier, then adjust for corpus size, iteration, tools, and agent count.")


def add_pricing_page(doc):
    doc.add_page_break()
    add_section_kicker(doc, "Foundations")
    add_heading(doc, "How to read the numbers", 1)
    add_callout(
        doc,
        "ONE FORMULA",
        "Estimated token cost = (input tokens ÷ 1,000,000 × input price) + "
        "(output tokens ÷ 1,000,000 × output price). For agents, add every agent and coordinator call.",
    )
    add_label_para(doc, "Cumulative input", "all prompts, system instructions, tool schemas, retrieved text, images represented as tokens, prior conversation history, and tool results sent back to a model.")
    add_label_para(doc, "Cumulative output", "visible responses plus billable reasoning/thinking and tool-call arguments. It is not just the final deliverable length.")
    add_label_para(doc, "Repeated context", "counts again when resent. A 100K-token project corpus read ten times can create roughly 1M input tokens even though the unique corpus is only 100K.")
    add_label_para(doc, "K / M", "K means 1,000 tokens; M means 1,000,000 tokens.")

    add_heading(doc, "Current standard API prices used in this guide", 1)
    add_caption(doc, "USD per 1M tokens. Standard first-party API rates; cached-input rates shown for reference. Batch, fast/priority, data residency, and tool-specific charges are excluded.")
    rows = [
        ["OpenAI", "GPT-5.6 Luna", "$0.20", "$0.02", "$1.20", "Bounded, high-volume, cost-sensitive work"],
        ["OpenAI", "GPT-5.6 Terra", "$2.00", "$0.20", "$12.00", "Balanced professional work and routine coding"],
        ["OpenAI", "GPT-5.6 Sol", "$5.00", "$0.50", "$30.00", "Complex reasoning, coding, and long-horizon work"],
        ["Anthropic", "Claude Haiku 4.5", "$1.00", "$0.10", "$5.00", "Fast, economical, bounded workloads"],
        ["Anthropic", "Claude Sonnet 5", "$2.00", "$0.20", "$10.00", "Coding, agents, analysis, and production synthesis"],
        ["Anthropic", "Claude Opus 5", "$5.00", "$0.50", "$25.00", "Complex agentic coding and enterprise work"],
        ["Anthropic", "Claude Fable 5", "$10.00", "$1.00", "$50.00", "Highest-capability long-running agents; escalate selectively"],
    ]
    add_simple_table(
        doc,
        ["Provider", "Model", "Input", "Cached", "Output", "Use it for"],
        rows,
        [1050, 1750, 850, 850, 850, 4010],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=8.7,
    )
    add_callout(
        doc,
        "MODEL ROUTING PRINCIPLE",
        "Start with the lowest tier that reliably passes your quality bar. Use Luna/Haiku for bounded steps, "
        "Terra/Sonnet for most production work, and Sol/Opus for ambiguity, difficult debugging, cross-domain synthesis, "
        "or long-horizon autonomy. Reserve Fable for cases where the highest capability is worth roughly 2× Opus token cost.",
    )


EVERYDAY_ROWS = [
    ("Ask a question", "Answer a bounded question from a short prompt; no external search.", 1_000, 4_000, 200, 1_000, "Luna", "Haiku 4.5"),
    ("Draft or rewrite", "Write an email, short brief, summary, or polished response with light context.", 2_000, 10_000, 500, 2_000, "Luna", "Haiku 4.5"),
    ("Do one Slack task", "Read a short thread, draft/send a message, or update one item through a connector.", 8_000, 40_000, 1_000, 5_000, "Luna", "Haiku 4.5"),
    ("Search Slack + synthesize", "Search a project history, read multiple threads, reconcile decisions, and produce a brief.", 30_000, 200_000, 3_000, 15_000, "Terra", "Sonnet 5"),
    ("Search a Figma file + synthesize", "Inspect frames/components, read annotations and imagery, compare patterns, and summarize findings.", 40_000, 250_000, 5_000, 25_000, "Terra", "Sonnet 5"),
    ("Summarize a document or meeting", "Read one long document or transcript and produce themes, decisions, and next steps.", 20_000, 120_000, 2_000, 10_000, "Terra", "Sonnet 5"),
]

KNOWLEDGE_ROWS = [
    ("Research brief with sources", "Search, retrieve, compare, cite, and synthesize a focused question across several sources.", 50_000, 300_000, 5_000, 25_000, "Terra", "Sonnet 5"),
    ("Spreadsheet / data analysis", "Inspect a workbook or dataset, run calculations, check anomalies, and explain the result.", 75_000, 400_000, 10_000, 50_000, "Terra", "Sonnet 5"),
    ("Professional doc or deck", "Research, structure, draft, format, revise, and visually verify a shareable artifact.", 75_000, 400_000, 10_000, 60_000, "Terra", "Sonnet 5"),
    ("Small code change", "Read relevant files, plan, edit, run tests, inspect failures, and revise.", 50_000, 250_000, 10_000, 50_000, "Terra", "Sonnet 5"),
    ("Debug or review a PR", "Trace behavior across files/logs, form hypotheses, test them, and report or implement a fix.", 150_000, 800_000, 25_000, 120_000, "Sol", "Opus 5"),
]

HEAVY_ROWS = [
    ("Deep knowledge work", "Spend 2–6 hours reading a corpus, developing hypotheses, testing counterarguments, and writing a decision-grade synthesis.", 800_000, 4_000_000, 150_000, 800_000, "Sol", "Opus 5"),
    ("Heavy coding day", "Explore a codebase, implement a substantial feature or refactor, test repeatedly, and document the result.", 1_000_000, 6_000_000, 200_000, 1_000_000, "Sol", "Opus 5"),
    ("MVP app from zero", "Clarify scope, scaffold, design UI/data flows, implement, integrate, test, polish, and package a deployable MVP.", 3_000_000, 15_000_000, 500_000, 2_500_000, "Sol", "Opus 5"),
]


def add_matrix_pages(doc):
    doc.add_page_break()
    add_section_kicker(doc, "Use-case matrix")
    add_heading(doc, "Everyday and connected work", 1)
    add_para(doc, "These are completed-task estimates, not single-response limits. Connector and server-side tool fees are separate.")
    add_matrix(doc, EVERYDAY_ROWS)
    add_caption(doc, "Figma note: screenshots, image representations, node metadata, and repeated inspection can dominate input. Slack note: a search across channels can become a research task even if the final output is only one page.")
    add_heading(doc, "How to place an unfamiliar task", 2)
    add_label_para(doc, "Bounded", "one clear instruction, little context, one output → start in the first three rows.")
    add_label_para(doc, "Synthesis", "many artifacts must be compared or reconciled → use the Figma/Slack synthesis or research rows.")
    add_label_para(doc, "Agentic", "the model chooses tools, loops after failures, edits files, and verifies work → use the coding or deep-work rows.")

    doc.add_page_break()
    add_section_kicker(doc, "Use-case matrix")
    add_heading(doc, "Knowledge production and software work", 1)
    add_matrix(doc, KNOWLEDGE_ROWS)
    add_caption(doc, "The upper end assumes several model/tool turns and at least one revision cycle. If the task is one-pass and tightly scoped, use the lower end.")

    add_heading(doc, "Deep and heavy work", 1)
    add_matrix(doc, HEAVY_ROWS)
    add_caption(doc, "“MVP from zero” means a modest deployable product, not a production system with security review, migration, observability, compliance, operations, and months of product iteration.")
    add_callout(
        doc,
        "WHY OUTPUT COST MATTERS",
        "Frontier-model output is 5–6× the price of input in these rate cards. Long reasoning traces, code generation, "
        "verbose agents, and repeated rewritten deliverables can move cost more than adding another source document.",
    )


def add_deep_agents_page(doc):
    doc.add_page_break()
    add_section_kicker(doc, "Deep work")
    add_heading(doc, "What “deep knowledge work” actually includes", 1)
    add_label_para(doc, "Frame", "define the decision, evidence standard, audience, uncertainties, and what would change the conclusion.")
    add_label_para(doc, "Retrieve", "search multiple corpora—documents, Slack, Figma, code, web sources, datasets—and preserve provenance.")
    add_label_para(doc, "Analyze", "build themes and models, compare alternatives, quantify where possible, and distinguish facts from inference.")
    add_label_para(doc, "Challenge", "look for contradictory evidence, missing stakeholders, weak causal claims, security/privacy issues, and false confidence.")
    add_label_para(doc, "Synthesize", "produce a decision-grade narrative, matrix, recommendation, or plan and iterate against the audience’s likely questions.")
    add_callout(
        doc,
        "DEEP-DEEP RANGE",
        "A single expert agent doing a half-day synthesis is roughly 0.8–4M input and 0.15–0.8M output. "
        "A full-day, cross-domain, decision-critical inquiry can reach 1.5–10M input and 0.25–1.5M output—about "
        "$15–$95 on OpenAI Sol or $14–$88 on Anthropic Opus 5 before tool fees.",
    )

    add_heading(doc, "Working with 1, 2, 4, or 8 agents", 1)
    add_caption(doc, "Illustrative heavy-work baseline: one agent uses 3M input + 0.5M output. Agent counts assume useful specialization and controlled handoffs; they are not a promise of speedup.")
    agent_rows = [
        ["1", "One owner: research → build → verify", ["3M in", "0.5M out"], ["OpenAI $30", "Anthropic $27.5"], "Not applicable"],
        ["2", "Coordinator + one specialist", ["5M in", "0.8M out", "≈1.7×"], ["OpenAI $49", "Anthropic $45"], ["OpenAI $27", "Anthropic $24.8"]],
        ["4", "Coordinator + 3 parallel specialists", ["9M in", "1.5M out", "≈3.0×"], ["OpenAI $90", "Anthropic $82.5"], ["OpenAI $49.5", "Anthropic $45.4"]],
        ["8", "Coordinator + 7 narrow workstreams", ["16M in", "2.6M out", "≈5.3×"], ["OpenAI $158", "Anthropic $145"], ["OpenAI $86.9", "Anthropic $79.8"]],
    ]
    add_simple_table(
        doc,
        ["Agents", "Useful topology", "Example cumulative use", "All frontier", "Mixed fleet"],
        agent_rows,
        [800, 2700, 1600, 2010, 2250],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.6,
    )
    add_caption(doc, "All-frontier = OpenAI Sol or Anthropic Opus 5 for every token. Mixed fleet = 25% of tokens on Sol/Opus and 75% on Terra/Sonnet. Real routing will differ by role and quality requirements.")
    doc.add_page_break()
    add_heading(doc, "When more agents help—and when they do not", 2)
    add_label_para(doc, "Good split", "independent source sets, separate technical subsystems, distinct hypotheses, or explicit researcher / builder / critic / verifier roles.")
    add_label_para(doc, "Poor split", "every agent rereads the same large corpus, edits the same files, or produces overlapping prose that the coordinator must reconcile.")
    add_label_para(doc, "Rule of thumb", "2 agents often cost 1.6–2.0× a strong single-agent run; 4 agents about 2.7–3.5×; 8 agents about 4.5–6.5×. Badly scoped swarms can approach or exceed linear growth.")


def add_local_example_page(doc):
    add_section_kicker(doc, "Local calibration")
    add_heading(doc, "A project already in this workspace", 1)
    add_para(doc, "The local public-GitHub exposure audit is a useful example of deep research and knowledge production—not a greenfield app build.")
    rows = [
        ["Scope", "Public-only account/repository discovery, deterministic scanning, evidence review, classification, and reporting"],
        ["Observed artifacts", "3 Python utilities totaling 985 lines; about 2.28 MB of reports, review packets, and JSON evidence in the inspected workspace"],
        ["Work volume", "160 person-linked public repositories plus a 128-repository IDEO organization pass represented in the finished overview"],
        ["Model-heavy steps", "identity/evidence synthesis, false-positive classification, cross-source reasoning, report drafting, revision, and provenance checks"],
        ["Mostly non-token steps", "GitHub/API retrieval, deterministic scans, local parsing, file generation, and shell execution"],
    ]
    add_simple_table(doc, ["Dimension", "What was observed"], rows, [1900, 7460],
                     [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.2)
    add_callout(
        doc,
        "PLANNING ESTIMATE—NOT MEASURED BILLING",
        "Based on artifact volume and the repeated synthesis/revision pattern, a comparable API workflow could reasonably budget "
        "2–8M cumulative input and 0.15–0.6M output. That is about $14.5–$58 on OpenAI Sol or $13.8–$55 on Anthropic Opus 5. "
        "A four-agent version could be higher if each specialist rereads the same evidence. Actual historical usage logs were not present, so this is an inference, not a reconstruction.",
    )
    doc.add_page_break()
    add_heading(doc, "Why the files do not equal the tokens", 2)
    add_label_para(doc, "Unique bytes understate use", "the same evidence may be retrieved, filtered, quoted, and resubmitted during several reasoning and revision turns.")
    add_label_para(doc, "Unique bytes can also overstate use", "deterministic scripts can scan megabytes or repositories without sending all content to a model.")
    add_label_para(doc, "The practical answer", "measure input, cached input, output, tool calls, retries, and model per call in production; use this guide only before telemetry exists.")

    add_heading(doc, "Fast planning worksheet", 1)
    plan_rows = [
        ["1. Deliverable", "Answer, Slack post, synthesis, report, code change, prototype, or deployable app"],
        ["2. Corpus", "How many files/frames/threads/repos? Approximate unique tokens or bytes"],
        ["3. Loops", "How many search, tool, test, critique, and revision cycles?"],
        ["4. Quality tier", "Bounded/fast, balanced production, or frontier/deep"],
        ["5. Parallelism", "1, 2, 4, or 8 agents—and a distinct role for each"],
        ["6. Non-token cost", "Search/tool fees, connector fees, image generation, code execution, hosting, storage, and human review"],
        ["7. Contingency", "Add 25–50% for ambiguous scope; 2× for first-of-kind work without stable evals"],
    ]
    add_simple_table(doc, ["Input", "Planning question"], plan_rows, [1800, 7560],
                     [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.2)


def add_assumptions_sources(doc):
    add_section_kicker(doc, "Guardrails")
    add_heading(doc, "Assumptions that materially change the estimate", 1)
    add_label_para(doc, "API-equivalent only", "ChatGPT, Codex, Claude, and Claude Code subscriptions/credits do not map one-to-one to the dollar estimates here.")
    add_label_para(doc, "Short-context OpenAI pricing", "the matrix assumes individual requests stay in the short-context tier. Cumulative task usage may be millions of tokens across many calls. A single request above the published threshold can use higher long-context rates.")
    add_label_para(doc, "Anthropic tokenization", "Anthropic notes that Claude 4.7-and-later tokenizers can produce about 30% more tokens than earlier Claude tokenizers for the same text; exact cross-provider token counts will differ.")
    add_label_para(doc, "Caching", "the matrix uses uncached input. Stable, repeated prefixes can be much cheaper when prompt caching is correctly implemented; new or changing context is not a cache hit.")
    add_label_para(doc, "Tools", "server-side web search and other metered tools can add charges. Client-side connectors still add tokenized tool definitions, arguments, and results even when the connector itself has no per-call fee.")
    add_label_para(doc, "Retries and failures", "broken tests, permission errors, weak search results, tool timeouts, and human revisions are normal sources of additional calls.")
    add_label_para(doc, "Security and privacy", "do not optimize solely for cost. Data handling, access controls, retention, provenance, and human review may require a different provider, model, or architecture.")

    add_heading(doc, "Recommended operating practice", 1)
    add_label_para(doc, "Estimate", "choose a use-case band and add a documented contingency.")
    add_label_para(doc, "Instrument", "log provider, model, input, cached input, output, reasoning/thinking, tool calls, latency, retries, and task outcome.")
    add_label_para(doc, "Evaluate", "compare quality and cost on representative tasks, not toy prompts.")
    add_label_para(doc, "Route", "use frontier models for the decisions and hard synthesis; use balanced or efficient models for retrieval, formatting, classification, and bounded subagents.")
    add_label_para(doc, "Review", "look at cost per successful task and human time saved, not cost per token in isolation.")

    add_heading(doc, "Official sources", 1)
    sources = [
        ("OpenAI API pricing", "https://developers.openai.com/api/docs/pricing", "standard, cached, long-context, batch/flex/fast pricing"),
        ("OpenAI model catalog", "https://developers.openai.com/api/docs/models", "current model positioning and selection guidance"),
        ("Anthropic pricing", "https://platform.claude.com/docs/en/about-claude/pricing", "standard, cache, batch, fast, residency, and tool-use pricing"),
        ("Anthropic model overview", "https://platform.claude.com/docs/en/about-claude/models/overview", "model IDs, context, latency, pricing, and capabilities"),
        ("Anthropic model selection", "https://platform.claude.com/docs/en/about-claude/models/choosing-a-model", "use-case guidance for Haiku, Sonnet, Opus, and Fable"),
    ]
    for idx, (name, url, note) in enumerate(sources, 1):
        p = doc.add_paragraph()
        style_paragraph(p, 0, 6, 1.2)
        r = p.add_run(f"{idx}. ")
        set_font(r, size=10.5, bold=True, color=INK)
        add_hyperlink(p, name, url)
        r = p.add_run(f" — {note}. Accessed August 12, 2026.")
        set_font(r, size=10.5, color=DARK)

    add_callout(
        doc,
        "KEEP THIS GUIDE CURRENT",
        "Model names and prices change. Recheck the two official pricing pages before using this matrix for procurement, client pricing, or a high-volume production forecast.",
    )


def audit_document(doc):
    for section in doc.sections:
        assert abs(section.page_width.inches - 8.5) < 0.01
        assert abs(section.page_height.inches - 11.0) < 0.01
        assert abs(section.left_margin.inches - 1.0) < 0.01
        assert abs(section.right_margin.inches - 1.0) < 0.01
    for table in doc.tables:
        grid = table._tbl.tblGrid
        widths = [int(c.get(qn("w:w"))) for c in grid]
        assert sum(widths) == 9360, (sum(widths), widths)
        tblW = table._tbl.tblPr.find(qn("w:tblW"))
        assert tblW is not None and int(tblW.get(qn("w:w"))) == 9360
        tblInd = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tblInd is not None and int(tblInd.get(qn("w:w"))) == 120


def build():
    doc = Document()
    setup_styles(doc)
    setup_section(doc.sections[0])
    props = doc.core_properties
    props.title = "AI Project Effort, Token & Cost Guide"
    props.subject = "Planning matrix for AI task token usage and API-equivalent cost"
    props.author = "IDEO"
    props.keywords = "AI, tokens, cost, OpenAI, Anthropic, agents, coding, knowledge work"

    add_cover(doc)
    add_pricing_page(doc)
    add_matrix_pages(doc)
    add_deep_agents_page(doc)
    add_local_example_page(doc)
    add_assumptions_sources(doc)
    audit_document(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
